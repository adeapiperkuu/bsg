"""Generate Workforce Agent comprehensive audit document (Word format)."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "Workforce-Agent-Comprehensive-Audit.docx"


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def add_h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def add_h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = val
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_title(doc, "Workforce & Capability Agent")
    add_title(doc, "Comprehensive Technical Audit")
    add_subtitle(doc, f"BSG Operations Tower — Audit Date: {date.today().strftime('%B %d, %Y')}")
    add_subtitle(doc, "Document language: English")
    doc.add_paragraph()

    # 1. Executive Summary
    add_h1(doc, "1. Executive Summary")
    add_para(
        doc,
        "This document provides a full technical audit of the Workforce & Capability Agent "
        "implementation across backend services, API layer, frontend UI, testing, security, "
        "and operational risks. The audit was performed against the current codebase in the "
        "BSG repository.",
    )
    add_bullets(
        doc,
        [
            "The active Workforce Agent is a deterministic, evidence-backed Q&A engine — it does not call an LLM for answer generation.",
            "Questions are routed through POST /api/v1/agent-queries with agent_name workforce_capability_agent.",
            "Answers are built from project-scoped database evidence: teams, utilization snapshots, skill matrix, training gaps, capability gaps, and risk alerts.",
            "Individual annotator names are never exposed; only aggregated counts are used.",
            "A legacy LLM-based implementation exists under app/agents/workforce/ but is not wired to the API.",
            "The frontend provides a chat UI integrated into the Workforce page with role-based access control.",
            "Evidence UI components exist but are not mounted in the live chat interface.",
            "Backend test suite: 44 test functions. Frontend: 32 Vitest test cases across 4 files.",
        ],
    )

    add_h2(doc, "1.1 Overall Assessment")
    add_table(
        doc,
        ["Dimension", "Rating", "Summary"],
        [
            ["Backend maturity", "High", "Well-structured deterministic service with privacy guardrails and extensive tests."],
            ["Frontend maturity", "Medium-High", "Clean component architecture; evidence UI built but not integrated."],
            ["Security & privacy", "High", "CLIENT role blocked; annotator names excluded; project scoping enforced."],
            ["Test coverage", "Medium-High", "Strong core flows; gaps in several intents, roles, and edge cases."],
            ["Architecture risk", "Medium", "Dual implementation confusion; keyword routing brittleness; API field mismatches."],
            ["Production readiness", "Medium-High", "Suitable for internal MVP; follow-ups needed for evidence UX and persistence."],
        ],
    )

    # 2. Scope
    add_h1(doc, "2. Audit Scope and Methodology")
    add_h2(doc, "2.1 In Scope")
    add_bullets(
        doc,
        [
            "backend/app/services/workforce_agent.py — active agent service (~906 lines)",
            "backend/app/api/routes/agents.py — API routing",
            "backend/app/services/agent_queries.py — agent registry",
            "backend/app/agents/workforce/ — legacy LLM implementation",
            "backend/tests/test_workforce_agent.py — backend tests",
            "frontend/src/components/bsg/workforce/agent/ — 18 source files",
            "frontend/src/routes/workforce.tsx — page integration",
            "Supporting services: workforce.py, workforce_skills.py, workforce_training.py, workforce_gaps.py, scoping.py, evidence.py",
        ],
    )
    add_h2(doc, "2.2 Out of Scope")
    add_bullets(
        doc,
        [
            "Workforce CRUD dashboards (teams, annotators, utilization charts) except where they feed agent evidence.",
            "Inter-agent signal consumer (skill_gap_consumer.py) except as architectural context.",
            "Infrastructure, deployment, and monitoring configuration.",
            "Performance benchmarking under production load.",
        ],
    )
    add_h2(doc, "2.3 Methodology")
    add_para(
        doc,
        "The audit combined static code review, dependency tracing, test inventory, "
        "comparison against the Developer 3 Roadmap document, and cross-layer contract analysis "
        "between backend answer format and frontend parsing logic.",
    )

    # 3. Product Context
    add_h1(doc, "3. Product Context")
    add_para(
        doc,
        "The Workforce & Capability Agent answers the question: Do we have the right workforce "
        "capability to deliver the work? It complements Delivery Performance (on-track?) and "
        "Quality Intelligence (quality stable?) agents.",
    )
    add_h2(doc, "3.1 Target Users")
    add_table(
        doc,
        ["Role", "App Role", "Access to Agent"],
        [
            ["Delivery Manager", "delivery_manager", "Full access"],
            ["BSG Leadership", "bsg_leadership", "Full access"],
            ["Super Admin", "super_admin", "Full access"],
            ["Client", "client", "Blocked (403)"],
        ],
    )
    add_h2(doc, "3.2 Supported Business Questions")
    add_bullets(
        doc,
        [
            "Which teams are overloaded or underutilized?",
            "Do we have enough SME coverage for this project?",
            "Which skills are missing or have low coverage?",
            "What training gaps exist (mandatory incomplete, expired/failed)?",
            "What certification gaps exist (expired, pending review)?",
            "What are the open capability gaps and their severity?",
            "What workforce recommendations exist from risk mitigation?",
            "Overall workforce capability summary for a project.",
        ],
    )

    # 4. Architecture
    add_h1(doc, "4. System Architecture")
    add_h2(doc, "4.1 End-to-End Request Flow")
    add_para(doc, "1. User submits a question in WorkforceAgentChat (frontend).")
    add_para(doc, "2. useWorkforceAgentChat calls createAgentQuery with agent_name workforce_capability_agent and project_id.")
    add_para(doc, "3. POST /api/v1/agent-queries routes to answer_workforce_query() when agent matches.")
    add_para(doc, "4. Service checks RBAC (can_read_annotators), classifies question scope, gathers evidence, builds answer.")
    add_para(doc, "5. AgentQuery and AgentQueryEvidenceLink rows are persisted.")
    add_para(doc, "6. Response returned; frontend formats answer_text via formatWorkforceAnswer().")

    add_h2(doc, "4.2 Dual Implementation (Critical Finding)")
    add_table(
        doc,
        ["Aspect", "Active (services/workforce_agent.py)", "Legacy (agents/workforce/query_handler.py)"],
        [
            ["API wired", "Yes", "No (lazy export only)"],
            ["Answer generation", "Deterministic templates", "LLM via LLMClient.generate_structured()"],
            ["Scoping", "Project-scoped", "Organization-scoped"],
            ["Utilization table", "UtilizationSnapshot (team-level)", "WorkforceUtilizationSnapshot"],
            ["Prompts", "None", "WORKFORCE_SYSTEM_PROMPT, WORKFORCE_USER_TEMPLATE"],
            ["Fallback", "N/A", "Pre-computed analysis_summary if LLM fails"],
        ],
    )
    add_para(
        doc,
        "Risk: Two functions share the name answer_workforce_query. Accidental re-wiring to the legacy "
        "handler would change behavior, data sources, and privacy characteristics.",
        bold=True,
    )

    # 5. Backend
    add_h1(doc, "5. Backend Audit")
    add_h2(doc, "5.1 Entry Point: answer_workforce_query()")
    add_bullets(
        doc,
        [
            "RBAC gate: raises ApiError(403) if can_read_annotators(current_user) is false.",
            "Out-of-scope redirect: classify_workforce_question() matches keywords against 4 other agents.",
            "Missing project_id: returns PROJECT_REQUIRED_MESSAGE (HTTP 200, not an error).",
            "In-scope with project: get_visible_project() → gather_workforce_evidence() → build_workforce_answer() → rank_evidence_for_intent().",
            "Persists AgentQuery with model_used from settings.llm_model (despite no LLM usage).",
            "Records latency_ms from perf_counter().",
        ],
    )

    add_h2(doc, "5.2 Out-of-Scope Redirect Rules")
    add_para(doc, "First-match-wins keyword routing redirects to owning agents:")
    add_table(
        doc,
        ["Target Agent", "Example Keywords"],
        [
            ["Delivery Performance Agent", "delivery confidence, milestone confidence, slippage, will we hit the deadline"],
            ["Quality Intelligence Agent", "quality drift, error rate, defect rate, rework rate, gold set"],
            ["Client Interaction Agent", "client email, csat, client satisfaction, client feedback"],
            ["Operational Knowledge Agent", "sop, standard operating procedure, policy document, knowledge base"],
        ],
    )

    add_h2(doc, "5.3 In-Scope Intent Detection")
    add_para(doc, "detect_workforce_intent() uses keyword matching (first match wins):")
    add_table(
        doc,
        ["Intent", "Keywords", "Answer Focus"],
        [
            ["sme", "sme, subject matter expert", "SME count, percentage, coverage assessment"],
            ["utilization", "overload, underload, utilization, at capacity", "Overloaded/underloaded teams, stale data warnings"],
            ["capacity", "headcount, annotator, staff, resources", "Active annotator counts, team structure"],
            ["skills", "skill, competenc", "Skill matrix, low-coverage skills"],
            ["training", "training", "Training gap counts, mandatory incomplete"],
            ["certification", "certif", "Expired certs, pending reviews"],
            ["capability_gaps", "capability gap, biggest gap", "Open gaps, high/critical counts"],
            ["recommendations", "recommend", "Mitigation recommendations from workforce risks"],
            ["summary", "summar, overall, overview, status (default)", "Generic project workforce summary"],
        ],
    )

    add_h2(doc, "5.4 Evidence Gathering")
    add_para(doc, "gather_workforce_evidence() queries project-scoped data:")
    add_table(
        doc,
        ["Source Table", "Data Collected", "Cited in Evidence?"],
        [
            ["teams", "Active teams for project", "Yes (up to 3)"],
            ["annotators", "Aggregated counts only (active, SME)", "No (privacy by design)"],
            ["utilization_snapshots", "Latest team-level snapshot per team", "Yes (up to 3)"],
            ["project_skill_requirements", "Requirement count + samples", "Yes (up to 3)"],
            ["training_programs / certifications", "Via build_project_training_gaps()", "Yes (up to 5, deduplicated)"],
            ["capability_gaps", "Open/acknowledged gaps", "Yes (up to 5)"],
            ["risk_alerts", "WORKFORCE_IMBALANCE type only", "Yes (up to 3)"],
            ["mitigation_recommendations", "Linked to workforce risk alerts", "Yes (up to 3)"],
        ],
    )

    add_h2(doc, "5.5 Thresholds and Constants")
    add_table(
        doc,
        ["Constant", "Value", "Purpose"],
        [
            ["UTILIZATION_OVERLOAD_THRESHOLD", "85%", "Team considered overloaded"],
            ["UTILIZATION_UNDERLOAD_THRESHOLD", "60%", "Team considered underutilized"],
            ["UTILIZATION_STALE_DAYS", "14", "Snapshots older than this flagged as stale"],
            ["WORKFORCE_CHAT_MAX_MESSAGE_LENGTH (frontend)", "2000", "Client-side question length limit"],
        ],
    )

    add_h2(doc, "5.6 Confidence Scoring")
    add_para(
        doc,
        "Confidence (High / Medium / Low) is embedded in answer_text footer as 'Confidence: High.' "
        "It is NOT written to retrieval_params JSONB. Therefore AgentQueryRead.confidence_level "
        "returns null for workforce queries despite the frontend attempting to read it.",
    )
    add_para(doc, "Confidence rules vary by intent — example for utilization:")
    add_bullets(
        doc,
        [
            "High: Fresh utilization data with overloaded or underloaded teams identified.",
            "Medium: Stale utilization snapshots (>14 days) or mixed signals.",
            "Low: No utilization snapshots available.",
        ],
    )

    add_h2(doc, "5.7 Answer Template Structure")
    add_para(doc, "All successful answers include a footer with:")
    add_bullets(
        doc,
        [
            "Grounded in N workforce record(s). — count of evidence items",
            "Confidence: High|Medium|Low.",
            "Optional Caution: lines for stale data or limitations",
            "Optional Next step: or Recommendation: prefixes for follow-up actions",
        ],
    )

    add_h2(doc, "5.8 Dependency Services")
    add_table(
        doc,
        ["Service", "Function", "Role in Agent"],
        [
            ["scoping.py", "get_visible_project()", "Enforces org/role project access"],
            ["workforce.py", "can_read_annotators()", "RBAC gate for agent access"],
            ["workforce_skills.py", "build_project_skill_matrix()", "Skill coverage analysis"],
            ["workforce_training.py", "build_project_training_gaps()", "Training/certification gaps"],
            ["workforce_gaps.py", "Threshold constants, OPEN_GAP_STATUSES", "Gap status filtering"],
            ["evidence.py", "EvidenceInput dataclass", "Pre-persistence evidence shape"],
        ],
    )

    # 6. API
    add_h1(doc, "6. API Contract Audit")
    add_h2(doc, "6.1 POST /api/v1/agent-queries")
    add_para(doc, "Request body (AgentQueryCreate):")
    add_bullets(
        doc,
        [
            "agent_name: string — must be workforce_capability_agent",
            "project_id: UUID | null — required for in-scope answers",
            "query_text: string (min 1 char)",
            "filters: object (optional) — accepted but IGNORED by workforce agent",
        ],
    )
    add_para(doc, "Response (AgentQueryRead) — workforce-relevant fields:")
    add_table(
        doc,
        ["Field", "Populated?", "Notes"],
        [
            ["answer_text", "Yes", "Primary answer content with embedded confidence"],
            ["evidence_links", "Yes", "0–N links to source tables"],
            ["model_used", "Yes", "Set to settings.llm_model despite no LLM call"],
            ["latency_ms", "Yes", "Request processing time"],
            ["confidence_level", "No (null)", "Not set in retrieval_params"],
            ["insufficient_evidence", "No (false)", "Not set for workforce agent"],
            ["related_records", "No (empty)", "Unused"],
            ["source_agents_used", "No (empty)", "Unused"],
            ["retrieval_params", "No (null)", "Unused"],
        ],
    )

    add_h2(doc, "6.2 GET /api/v1/agent-queries and GET /api/v1/agent-queries/{id}")
    add_bullets(
        doc,
        [
            "List recent queries across ALL agents (no agent_name filter).",
            "Access: client sees own queries; delivery_manager sees org; bsg_leadership/super_admin see all orgs.",
            "No dedicated workforce-only history endpoint.",
            "Frontend does NOT use these endpoints — chat history is in-memory only.",
        ],
    )

    add_h2(doc, "6.3 Error Responses")
    add_table(
        doc,
        ["Condition", "HTTP Status", "Error Code"],
        [
            ["Unknown agent name", "400", "VALIDATION_ERROR"],
            ["CLIENT role", "403", "FORBIDDEN"],
            ["Cross-org / inaccessible project", "403", "FORBIDDEN"],
            ["Project not found", "404", "NOT_FOUND"],
            ["Missing authentication", "401", "AUTH_REQUIRED"],
            ["Invalid CSRF (cookie auth)", "403", "CSRF_FAILED"],
        ],
    )

    # 7. Security
    add_h1(doc, "7. Security and Privacy Audit")
    add_h2(doc, "7.1 Access Control")
    add_bullets(
        doc,
        [
            "ANNOTATOR_READ_ROLES: delivery_manager, bsg_leadership, super_admin.",
            "CLIENT role blocked at service entry with 403 FORBIDDEN.",
            "Frontend mirrors this via canReadInternalWorkforce prop — shows restricted placeholder for clients.",
            "Project scoping via get_visible_project() prevents cross-org data access.",
        ],
    )
    add_h2(doc, "7.2 Privacy Protections")
    add_bullets(
        doc,
        [
            "Annotator full_name never appears in answers or evidence descriptions.",
            "Only aggregated counts: active annotators, SME count, SME percentage.",
            "annotators table is queried for counts but never cited as evidence source.",
            "Backend tests explicitly verify name exclusion.",
        ],
    )
    add_h2(doc, "7.3 Security Gaps")
    add_table(
        doc,
        ["Gap", "Severity", "Description"],
        [
            ["No rate limiting on agent-queries POST", "Low", "Delivery chat has rate limits; workforce agent does not"],
            ["Cross-org query listing for leadership", "Medium", "BSG_LEADERSHIP/SUPER_ADMIN can list all orgs' queries"],
            ["No X-BSG-User-Action header", "Low", "Knowledge agent requires this; workforce does not"],
            ["model_used misleading", "Low", "Suggests LLM usage in audit logs when none occurred"],
        ],
    )

    # 8. Frontend
    add_h1(doc, "8. Frontend Audit")
    add_h2(doc, "8.1 Component Hierarchy")
    add_para(doc, "WorkforcePage → WorkforceAgentSection → WorkforceAgentChat")
    add_para(doc, "WorkforceAgentChat contains:")
    add_bullets(
        doc,
        [
            "WorkforceAgentHistory — popover for in-memory session list",
            "WorkforceAgentSuggestions — 5 starter prompts (hidden after first message)",
            "WorkforceAgentMessage[] — chat bubbles with Bot avatar for agent",
            "WorkforceAgentAnswer — structured answer rendering",
            "TypingIndicator — 'Analyzing workforce data' loading state",
            "WorkforceAgentComposer — single-line input, Enter to send",
        ],
    )

    add_h2(doc, "8.2 State Management (useWorkforceAgentChat)")
    add_table(
        doc,
        ["State", "Type", "Persistence"],
        [
            ["sessions", "WorkforceChatSession[]", "In-memory only (lost on refresh)"],
            ["activeSessionId", "string", "In-memory"],
            ["input", "string", "In-memory"],
            ["asking", "boolean", "In-memory"],
            ["error", "string | null", "In-memory"],
        ],
    )
    add_para(doc, "Key behaviors:")
    add_bullets(
        doc,
        [
            "Optimistic user message on send; agent message with full AgentQueryRead on success.",
            "New chat reuses empty session for project if one exists.",
            "Project switch loads most recent session for that project or creates new one.",
            "Message length guard at 2000 chars (hook + composer maxLength).",
            "Uses useMutation(createAgentQuery) — no React Query cache for history.",
        ],
    )

    add_h2(doc, "8.3 Answer Formatting (format-workforce-answer.ts)")
    add_para(doc, "~295 lines. Parses backend answer_text into structured UI sections:")
    add_table(
        doc,
        ["Section", "Source", "Displayed?"],
        [
            ["Headline", "Heuristic from utilization/SME/skill patterns", "Yes"],
            ["Summary", "First key finding or second body line", "Yes"],
            ["Key findings", "Lines with prefixes (overloaded:, skill gap:, etc.)", "Yes"],
            ["Data freshness", "Stale utilization warnings reclassified from Caution", "Yes"],
            ["Caution", "Non-freshness warnings", "Yes"],
            ["Suggested next step", "Next step:/Recommendation: prefixes", "Yes"],
            ["Confidence badge", "Parsed from Confidence: line OR confidence_level field", "Yes"],
            ["Limited evidence badge", "insufficient_evidence flag", "Yes (when true)"],
            ["Grounded in N records", "Parsed but not displayed", "No"],
            ["Evidence cards", "evidence_links via WorkforceAgentEvidence", "No (component not mounted)"],
            ["model_used / latency_ms", "API response fields", "No"],
        ],
    )

    add_h2(doc, "8.4 Starter Questions")
    add_bullets(
        doc,
        [
            "Which teams are overloaded?",
            "Do we have enough SME coverage?",
            "What are the biggest capability gaps?",
            "Which skills are missing for this project?",
            "Are training gaps creating risk?",
        ],
    )

    add_h2(doc, "8.5 Page Integration")
    add_bullets(
        doc,
        [
            "Located in workforce.tsx right column (lg:col-span-2), below By Region and Training Gaps cards.",
            "Receives resolvedProjectId from URL search param or first available project.",
            "Same project context as workforce dashboard KPIs, skill matrix, utilization chart.",
            "Agent does NOT read React Query workforce caches — each question is a fresh API call.",
        ],
    )

    add_h2(doc, "8.6 Comparison with Governance Agent")
    add_table(
        doc,
        ["Feature", "Workforce Agent", "Governance Agent"],
        [
            ["Evidence panel", "Built but not wired", "Integrated with EvidenceBadge"],
            ["Typewriter animation", "No", "Yes"],
            ["Copy / retry controls", "No", "Yes"],
            ["Session persistence", "In-memory only", "In-memory only"],
            ["Project scope selector", "Uses page project", "Includes 'all projects' option"],
            ["Toast on error", "Inline error only", "Toast notification"],
        ],
    )

    # 9. Testing
    add_h1(doc, "9. Testing and Quality Assurance")
    add_h2(doc, "9.1 Backend Tests (test_workforce_agent.py)")
    add_para(doc, "44 test functions covering:")
    add_table(
        doc,
        ["Area", "Coverage Level"],
        [
            ["Agent registration in SUPPORTED_AGENTS", "Full"],
            ["Out-of-scope redirects (all 4 agents)", "Full"],
            ["RBAC — CLIENT blocked (service + HTTP)", "Full"],
            ["Cross-org project access", "Full"],
            ["Privacy — no annotator names", "Full"],
            ["Intent answers (SME, utilization, skills, training, gaps, recommendations, summary)", "Good"],
            ["Confidence levels (utilization)", "Good"],
            ["Evidence ranking (SME, utilization)", "Partial"],
            ["Stale utilization warnings", "Full"],
            ["Certification intent with gaps", "Not tested"],
            ["Recommendations with data present", "Not tested"],
            ["Training with total > 0", "Not tested"],
            ["BSG_LEADERSHIP / SUPER_ADMIN access", "Not tested"],
            ["List/get agent-queries for workforce", "Minimal"],
            ["Legacy query_handler.py", "Not tested via API"],
        ],
    )

    add_h2(doc, "9.2 Frontend Tests (Vitest)")
    add_table(
        doc,
        ["File", "Test Cases", "Coverage"],
        [
            ["format-workforce-answer.test.ts", "10", "humanizeWording, headlines, stale data, next step"],
            ["WorkforceAgentAnswer.test.tsx", "12", "evidence-utils, evidence hidden assertion, answer rendering"],
            ["useWorkforceAgentChat.test.tsx", "5", "Send/receive, new chat, history, errors"],
            ["WorkforceAgentSection.test.tsx", "5", "Integration: suggestions, errors, history"],
        ],
    )
    add_para(doc, "Frontend gaps not tested:")
    add_bullets(
        doc,
        [
            "canReadInternalWorkforce=false (restricted placeholder)",
            "projectId=null / no-project state",
            "Project switching session restore",
            "insufficient_evidence badge display",
            "Message length limit enforcement",
            "Individual component tests for Composer, History, Suggestions, Chat, Message",
            "Accessibility (ARIA, keyboard navigation)",
            "E2E / real API integration",
        ],
    )

    # 10. Risks
    add_h1(doc, "10. Risks, Gaps, and Recommendations")
    add_h2(doc, "10.1 High Priority")
    add_table(
        doc,
        ["#", "Risk / Gap", "Impact", "Recommendation"],
        [
            ["H1", "Dual implementation (active vs legacy LLM)", "Accidental re-wiring changes behavior and data model", "Remove or clearly deprecate legacy handler; add import guard comment in agents.py"],
            ["H2", "confidence_level API field not populated", "Frontend structured confidence may be inconsistent", "Write confidence to retrieval_params or parse only from answer_text consistently"],
            ["H3", "Evidence UI not wired", "Users cannot verify 'why this answer' despite backend returning evidence_links", "Mount WorkforceAgentEvidence in WorkforceAgentAnswer behind expandable section"],
        ],
    )

    add_h2(doc, "10.2 Medium Priority")
    add_table(
        doc,
        ["#", "Risk / Gap", "Impact", "Recommendation"],
        [
            ["M1", "Keyword-based routing fragility", "False redirects or intent misclassification", "Add integration tests for edge-case phrasing; consider intent scoring"],
            ["M2", "has_data requires teams > 0", "Projects with gaps/training but no teams get insufficient-data", "Relax gate or provide partial answers"],
            ["M3", "No session persistence", "Chat history lost on refresh", "Persist to localStorage or use GET /agent-queries API"],
            ["M4", "Session switch race during ask", "Response may attach to wrong session", "Cancel in-flight request or lock session during ask"],
            ["M5", "Evidence sampling (first N only)", "Answer references aggregates beyond cited evidence", "Document limitation or increase sample size"],
            ["M6", "Cross-org query listing", "Leadership roles see all orgs' queries", "Add org filter or document as intended"],
        ],
    )

    add_h2(doc, "10.3 Low Priority")
    add_table(
        doc,
        ["#", "Risk / Gap", "Impact", "Recommendation"],
        [
            ["L1", "model_used misleading in audit trail", "Monitoring confusion", "Set to 'deterministic' or null for workforce agent"],
            ["L2", "No rate limiting", "Potential abuse", "Apply same limits as delivery chat"],
            ["L3", "Duplicate API helpers (createAgentQuery vs postAgentQuery)", "Maintenance drift", "Consolidate to single export"],
            ["L4", "filters field ignored", "API contract confusion", "Document as unused or implement"],
            ["L5", "Unused AppRole import in workforce_agent.py", "Lint noise", "Remove unused import"],
            ["L6", "groundedIn parsed but not shown", "Lost trust signal", "Display in answer footer"],
        ],
    )

    # 11. File Inventory
    add_h1(doc, "11. File Inventory")
    add_h2(doc, "11.1 Backend")
    add_table(
        doc,
        ["File", "Role", "Lines (approx.)"],
        [
            ["backend/app/services/workforce_agent.py", "Active agent service", "~906"],
            ["backend/app/api/routes/agents.py", "API endpoints", "—"],
            ["backend/app/services/agent_queries.py", "Agent registry", "—"],
            ["backend/tests/test_workforce_agent.py", "Test suite (44 tests)", "—"],
            ["backend/app/agents/workforce/query_handler.py", "Legacy LLM handler (unwired)", "—"],
            ["backend/app/agents/workforce/prompts.py", "Legacy prompts (unwired)", "—"],
            ["backend/app/agents/workforce/skill_gap_consumer.py", "Inter-agent signal consumer", "—"],
            ["backend/app/schemas/domain.py", "AgentQueryCreate/Read schemas", "—"],
            ["backend/app/services/evidence.py", "EvidenceInput dataclass", "—"],
            ["backend/app/services/scoping.py", "Project visibility", "—"],
            ["backend/app/services/workforce.py", "RBAC helpers", "—"],
            ["backend/app/services/workforce_skills.py", "Skill matrix builder", "—"],
            ["backend/app/services/workforce_training.py", "Training gap builder", "—"],
            ["backend/app/services/workforce_gaps.py", "Gap thresholds", "—"],
        ],
    )

    add_h2(doc, "11.2 Frontend")
    add_table(
        doc,
        ["File", "Role"],
        [
            ["WorkforceAgentSection.tsx", "Role gate + chat shell entry"],
            ["WorkforceAgentChat.tsx", "Card layout, scroll, header controls"],
            ["WorkforceAgentMessage.tsx", "Chat bubble layout"],
            ["WorkforceAgentAnswer.tsx", "Structured answer rendering"],
            ["WorkforceAgentEvidence.tsx", "Evidence cards (not mounted in live UI)"],
            ["WorkforceAgentComposer.tsx", "Input + send button"],
            ["WorkforceAgentHistory.tsx", "Session history popover"],
            ["WorkforceAgentSuggestions.tsx", "Starter prompt pills"],
            ["useWorkforceAgentChat.ts", "Chat state hook"],
            ["format-workforce-answer.ts", "Answer text parser (~295 lines)"],
            ["evidence-utils.ts", "Evidence label mapping and normalization"],
            ["session-utils.ts", "Session ID, title, time formatting"],
            ["types.ts", "WorkforceChatMessage, WorkforceChatSession"],
            ["constants.ts", "Agent name, max length, starter questions"],
            ["workforce.tsx", "Page integration (Phase 6 sidebar)"],
        ],
    )

    # 12. Appendix
    add_h1(doc, "12. Appendix")
    add_h2(doc, "12.1 Agent Constants")
    add_bullets(
        doc,
        [
            "WORKFORCE_AGENT_NAME = workforce_capability_agent",
            "UTILIZATION_STALE_DAYS = 14",
            "UTILIZATION_OVERLOAD_THRESHOLD = 85%",
            "UTILIZATION_UNDERLOAD_THRESHOLD = 60%",
            "WORKFORCE_CHAT_MAX_MESSAGE_LENGTH = 2000",
            "DEFAULT_EVIDENCE_PREVIEW_COUNT = 4",
        ],
    )

    add_h2(doc, "12.2 Evidence Source Table Labels (Frontend)")
    add_para(
        doc,
        "utilization_snapshots → Utilization snapshot; teams → Team; "
        "project_skill_requirements → Skill requirement; capability_gaps → Capability gap; "
        "risk_alerts → Risk alert; mitigation_recommendations → Recommendation; "
        "training_programs → Training program; certifications → Certification; "
        "projects → Project.",
    )

    add_h2(doc, "12.3 Document Revision History")
    add_table(
        doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["1.0", date.today().strftime("%Y-%m-%d"), "Automated audit", "Initial comprehensive audit"],
        ],
    )

    add_para(doc, "")
    p = doc.add_paragraph()
    run = p.add_run("— End of Document —")
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(OUTPUT))
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    main()
