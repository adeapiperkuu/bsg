"""Generate BSG Operations Tower build plan (Knowledge + Workforce agents)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = REPO_ROOT / "docs" / "BSG_Operations_Tower_Knowledge_Workforce_Build_Plan.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Operations Tower")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Technical Build Plan  ·  Knowledge Agent  ·  Workforce Agent\n"
        "Based on: BSG Ops Intelligence Agent v1.0  ·  AI Suite Build Plan"
    )
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
        set_cell_shading(hdr_cells[i], "E8EEF4")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    add_title_block(doc)

    # 1. What we are building
    add_heading(doc, "1. What We Are Building")
    doc.add_paragraph(
        "Operations Tower is BSG's AI-native operational intelligence platform — a role-scoped web "
        "cockpit for delivery managers, QA leads, operations managers, and BSG leadership. "
        "This document covers the current engineering workstream for two Phase 2 agents only:"
    )
    add_bullets(
        doc,
        [
            "Operational Knowledge Agent — RAG over SOPs, guides, training documents, and historical lessons; "
            "conversational process guidance with evidence citations.",
            "Workforce & Capability Agent — skills taxonomy, utilization, certifications, training gaps, "
            "capability gap detection, and workforce NL Q&A.",
        ],
    )
    doc.add_paragraph(
        "Four other agents in the original six-agent suite — Delivery Performance, Quality Intelligence, "
        "Project Governance, and Client Interaction — are not started. "
        "This document covers only the Knowledge and Workforce workstreams."
    )

    # 2. Scope in this repo
    add_heading(doc, "2. In Scope vs. Not Started")
    add_table(
        doc,
        ["Component", "Status in current repo", "Notes"],
        [
            ["Operational Knowledge Agent", "In active development", "Full library, ingestion, hybrid RAG, streaming Q&A"],
            ["Workforce & Capability Agent", "In active development", "Dashboards, skills matrix, gaps, agent Q&A"],
            ["Delivery Performance Agent", "Not started", "Planned Phase 1 — out of this document"],
            ["Quality Intelligence Agent", "Not started", "Planned Phase 1 — out of this document"],
            ["Project Governance Agent", "Not started", "Planned Phase 2 — out of this document"],
            ["Client Interaction Agent", "Not started", "Planned Phase 1 — out of this document"],
            ["Shared platform (auth, RLS, DB)", "Implemented", "Supports both agents today"],
        ],
        [1.6, 1.4, 2.5],
    )

    # 3. Tech stack
    add_heading(doc, "3. Technical Stack & Infrastructure")
    add_table(
        doc,
        ["Layer", "Technology", "Role in Knowledge / Workforce"],
        [
            ["Frontend", "React 18, TypeScript, Vite, TanStack Router & Query, Tailwind CSS 4, shadcn/ui, Recharts", "/knowledge and /workforce routes; SSE streaming chat; dashboard charts"],
            ["Backend", "Python 3.12, FastAPI, SQLAlchemy 2 (async), asyncpg, Pydantic v2, APScheduler", "REST + SSE APIs; background document processing; weekly jobs"],
            ["Database", "PostgreSQL 16 (Supabase), RLS, pgvector (HNSW)", "Tenant isolation; vector embeddings for Knowledge RAG"],
            ["Auth", "Supabase Auth, JWT, httpOnly cookies, CSRF on mutations", "Roles: delivery_manager, bsg_leadership, super_admin; clients excluded from workforce"],
            ["LLM", "OpenAI (gpt-4o-mini, gpt-4o, text-embedding-3-small)", "Knowledge ask/stream, grounding checks; Workforce NL Q&A via agent_queries"],
            ["File processing", "PyMuPDF, mammoth, openpyxl", "PDF/DOCX/TXT/MD/CSV ingestion for knowledge library"],
            ["Storage", "Local filesystem (KNOWLEDGE_UPLOAD_DIR) / Supabase-ready", "Document binaries and versions"],
            ["Repo layout", "Monorepo: frontend/, backend/, supabase/", "Migrations in supabase/migrations/"],
            ["Deployment", "Dockerfile (backend), run_dev_server.ps1; staging/prod TBD", "No CI/CD workflows in repo yet"],
        ],
        [1.1, 1.8, 2.6],
    )

    doc.add_paragraph(
        "Architecture pattern: three layers — (1) governed operational data in PostgreSQL with RLS, "
        "(2) deterministic analytics and dashboards, (3) conversational AI with evidence links on every answer. "
        "Knowledge uses dedicated /api/v1/knowledge/* endpoints; Workforce Q&A uses /api/v1/agent-queries "
        "with agent_name workforce_capability_agent."
    )

    add_heading(doc, "3.1 Infrastructure & Data Flow")
    add_bullets(
        doc,
        [
            "Request path: React UI → FastAPI (cookie JWT + CSRF) → SQLAlchemy async session → PostgreSQL with RLS context (org_id, role).",
            "Knowledge ingestion: multipart upload → local storage → background job (extract → chunk → embed via OpenAI) → pgvector HNSW index.",
            "Knowledge retrieval: hybrid vector + keyword search, query rewrite for follow-ups, answer grounding validation, confidence scoring.",
            "Workforce analytics: project-scoped teams and utilization snapshots; skill matrix computed from annotator_skills vs project_skill_requirements.",
            "Inter-agent readiness: skill_gap_consumer and lesson_log hooks exist; active when Quality agent is built.",
        ],
    )

    add_heading(doc, "3.2 Key Environment & Config")
    add_table(
        doc,
        ["Variable / setting", "Purpose"],
        [
            ["DATABASE_URL, SUPABASE_*", "PostgreSQL + Auth"],
            ["OPENAI_API_KEY / LLM_*", "Embeddings and chat completions"],
            ["KNOWLEDGE_UPLOAD_DIR", "Document binary storage"],
            ["ALLOWED_ORIGINS", "CORS for frontend BFF"],
            ["AUTH_COOKIE_SECURE", "Production cookie hardening"],
        ],
        [2.0, 3.5],
    )

    # 4. Modules
    add_heading(doc, "4. Modules — Knowledge Agent")
    add_table(
        doc,
        ["Module", "Location", "Capability", "Status"],
        [
            ["API routes", "backend/app/api/routes/knowledge.py", "Bootstrap, folders, documents CRUD, ask/stream, conversations, gaps, feedback, retrieval settings", "Done"],
            ["Orchestration", "backend/app/services/knowledge.py", "Upload pipeline, chunking, embedding, hybrid retrieval, grounding, gap detection", "Done"],
            ["LLM layer", "backend/app/services/llm/", "Structured answers, fast-path, client-safe mode, streaming", "Done"],
            ["Lesson write-back", "backend/app/agents/knowledge/lesson_log.py", "Quality alert resolve → knowledge_lessons", "Done"],
            ["DB schema", "supabase/migrations/20260623* – 20260702*", "Documents, chunks, embeddings, gaps, conversations, extraction metadata", "Done"],
            ["UI", "frontend/src/routes/knowledge.tsx", "Library browser, upload, streaming chat, history, gaps, settings", "Done"],
        ],
        [1.0, 1.5, 2.2, 0.8],
    )

    add_heading(doc, "4.1 Knowledge — API Surface")
    add_bullets(
        doc,
        [
            "GET /knowledge/bootstrap, /library-health — folder tree, permissions, gap counts.",
            "POST /knowledge/documents — upload; GET/PATCH/DELETE for metadata and lifecycle.",
            "POST /knowledge/ask, /knowledge/ask/stream — sync and SSE streaming Q&A.",
            "GET /knowledge/conversations — multi-turn history; POST /knowledge/feedback — thumbs up/down.",
            "GET/PATCH /knowledge/retrieval-settings — per-org hybrid retrieval tuning.",
        ],
    )

    add_heading(doc, "5. Modules — Workforce Agent")
    add_table(
        doc,
        ["Module", "Location", "Capability", "Status"],
        [
            ["API routes", "backend/app/api/routes/workforce.py", "Dashboard, teams, utilization, skills, certs, training, capability gaps", "Done"],
            ["Services", "backend/app/services/workforce*.py", "CRUD, gap detection, recommendations, privacy-safe agent Q&A", "Done"],
            ["Agent handler", "backend/app/services/workforce_agent.py", "Evidence-backed NL queries; no individual annotator names to clients", "Done"],
            ["Signal consumer", "backend/app/agents/workforce/skill_gap_consumer.py", "Consumes skill_gap signals from Quality (when Quality ships)", "Done"],
            ["DB schema", "supabase/migrations/202606251*", "skills, annotator_skills, utilization_snapshots, capability_gaps, training", "Done"],
            ["UI", "frontend/src/routes/workforce.tsx", "Utilization charts, skill matrix, gaps, training summary, agent panel", "Done"],
        ],
        [1.0, 1.5, 2.2, 0.8],
    )

    add_heading(doc, "5.1 Workforce — API Surface")
    add_bullets(
        doc,
        [
            "GET /workforce/dashboard, /skill-matrix, /sme-allocation — project KPIs and coverage.",
            "CRUD for teams, annotators, utilization snapshots, skills, certifications, training records.",
            "POST /workforce/capability-gaps/detect — automated gap detection with deduplication.",
            "POST /agent-queries (workforce_capability_agent) — NL Q&A with evidence links.",
        ],
    )

    # Milestones
    add_heading(doc, "6. Milestones & Dynamic Plan")
    doc.add_paragraph(
        "The dynamic plan below reflects actual repo progress (July 2026). "
        "Dates are indicative; scope adjusts as pilot feedback arrives."
    )
    add_table(
        doc,
        ["Milestone", "Target", "Deliverables", "Status"],
        [
            ["M0 — Platform foundation", "Weeks 0–2", "Monorepo, Supabase schema, RLS, auth cookies, agent_queries table", "Complete"],
            ["M1 — Workforce data model", "Weeks 3–5", "Teams, skills taxonomy, utilization snapshots, certifications, training", "Complete"],
            ["M2 — Workforce dashboards & Q&A", "Weeks 5–7", "Skill matrix, capability gaps, NL agent with evidence", "Complete"],
            ["M3 — Knowledge ingestion", "Weeks 5–7", "Upload, extraction, pgvector embeddings, reindex", "Complete"],
            ["M4 — Knowledge RAG & chat", "Weeks 7–9", "Hybrid retrieval, streaming ask, conversations, feedback", "Complete"],
            ["M5 — Knowledge library UX", "Weeks 9–10", "Folder library, gaps todos, retrieval settings, health KPIs", "Complete"],
            ["M6 — Hardening & pilot prep", "Weeks 10–12", "Tests, performance indexes, governance sign-off, sanitized pilot data", "In progress"],
            ["M7 — Production pilot", "Weeks 12–14", "Internal pilot on synthetic data; human review of all AI outputs", "Planned"],
        ],
        [1.3, 0.7, 2.5, 0.8],
    )

    add_heading(doc, "7. Dynamic Plan — Next 4 Weeks")
    add_bullets(
        doc,
        [
            "Knowledge: tighten extraction quality scoring, conversation history polish, retrieval tuning per org.",
            "Workforce: automated utilization ingestion (currently manual CRUD), leadership portfolio view.",
            "Shared: CI pipeline (GitHub Actions), Playwright tenant-isolation e2e, staging environment.",
            "Deferred until other agents start: inter-agent signals (Quality → Workforce skill_gap), lesson write-back from live alerts.",
        ],
    )

    # 6. Other agents
    add_heading(doc, "8. Other Agents — Not Started")
    doc.add_paragraph(
        "The following agents from the original AI Suite six-agent design are not part of the current "
        "engineering sprint. No frontend routes, backend handlers, or dedicated schemas are being built "
        "for them in this workstream:"
    )
    add_table(
        doc,
        ["Agent", "Planned phase", "Status"],
        [
            ["Delivery Performance Agent", "Phase 1", "Not started — milestone tracking, risk prediction, throughput forecasting"],
            ["Quality Intelligence Agent", "Phase 1", "Not started — drift detection, error taxonomy, RCA"],
            ["Client Interaction Agent", "Phase 1", "Not started — client summaries, confidence scores, approval workflow"],
            ["Project Governance Agent", "Phase 2", "Not started — charter, escalations, dependency tracking"],
        ],
        [2.0, 1.2, 2.3],
    )

    # Timeline summary
    add_heading(doc, "9. Timeline Summary")
    add_table(
        doc,
        ["Workstream", "Duration (est.)", "Team focus", "Key output"],
        [
            ["Shared infrastructure", "2 weeks", "1 backend + 1 devops", "Auth, RLS, monorepo, migrations"],
            ["Workforce Agent", "4–5 weeks", "1 backend + 1 frontend", "Dashboards, skills, gaps, Q&A"],
            ["Knowledge Agent", "5–6 weeks", "1 backend + 1 frontend + AI", "RAG library, ingestion, streaming chat"],
            ["Hardening & pilot", "2 weeks", "Full team", "Tests, security review, internal pilot"],
        ],
        [1.4, 1.0, 1.5, 2.6],
    )
    doc.add_paragraph(
        "Total elapsed for Knowledge + Workforce (parallel workstreams): approximately 10–12 weeks with "
        "a 3–4 person team. Remaining four agents would add an estimated 8–10 weeks if built sequentially "
        "after Phase 1 trio completion."
    )

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(
        f"Borek Solutions Group  ·  AI Engineering Team  ·  {date.today().strftime('%B %Y')}  ·  Confidential"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r.italic = True

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
