"""DOCX exporter with python-docx and minimal OOXML fallback."""

from __future__ import annotations

import io

from app.db.models import ReportInstance
from app.reports.exporters.base import ExportArtifact, safe_stem


def render(report: ReportInstance) -> ExportArtifact:
    try:
        from docx import Document

        document = Document()
        document.add_heading(report.title, level=0)
        for raw in (report.body_markdown or "").splitlines():
            line = raw.strip()
            if not line:
                document.add_paragraph()
            elif line.startswith("### "):
                document.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                document.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                document.add_heading(line[2:], level=1)
            elif line.startswith(("- ", "* ")):
                document.add_paragraph(line[2:], style="List Bullet")
            else:
                document.add_paragraph(line)
        buffer = io.BytesIO()
        document.save(buffer)
        content = buffer.getvalue()
    except ImportError:
        from app.agents.governance.services.charter_export import (
            CharterExportDocument,
            generate_charter_docx,
        )

        content = generate_charter_docx(
            CharterExportDocument(
                title=report.title,
                metadata=[
                    ("Template", f"{report.template_key}@{report.template_version}"),
                    ("Status", report.status),
                    ("Audience", report.audience),
                ],
                markdown=report.body_markdown or "",
            )
        )
    return (
        content,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"{safe_stem(report)}.docx",
    )
